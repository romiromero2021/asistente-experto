# -*- coding: utf-8 -*-
"""Untitled45.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1Dtcff-Vn7hs1iqrbPrXEvhL1pqgL-2dd
"""

import streamlit as st
import io
import os
import openai
from openai import OpenAI
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import base64

# Configuración de la página
st.set_page_config(
    page_title="Asistente Energético GPT",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilo CSS personalizado
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    .stProgress>div>div>div {
        background-color: #1f77b4;
    }
    </style>
""", unsafe_allow_html=True)

class AsistenteEnergetico:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key)
        self.contexto_sistema = """Eres un experto en temas energéticos, específicamente en:
        - Electrificación rural
        - Cocción de alimentos y uso de energía
        - Acceso a energía moderna
        - Políticas energéticas

        Tu tarea es analizar documentos sobre estos temas y responder preguntas de manera detallada y precisa,
        citando información específica de los documentos cuando sea relevante."""

    @st.cache_data
    def procesar_pdf(self, archivo_pdf):
        """Procesa el PDF y extrae el texto con caché"""
        try:
            pdf_reader = PyPDF2.PdfReader(archivo_pdf)
            texto = ""
            for pagina in pdf_reader.pages:
                texto += pagina.extract_text() + "\n"
            return texto
        except Exception as e:
            st.error(f"Error al procesar el PDF: {str(e)}")
            return None

    def generar_resumen(self, texto_documento, nombre_documento, progress_callback=None):
        """Genera un resumen estructurado del documento"""
        try:
            # Limitar el texto para el análisis inicial
            texto_inicial = texto_documento[:15000]

            if progress_callback:
                progress_callback(0.3, "Analizando contenido principal...")

            prompt_inicial = f"""
            Genera un resumen ejecutivo estructurado del siguiente documento: {nombre_documento}

            {texto_inicial}

            Estructura el resumen en:
            1. Objetivo principal (2-3 oraciones)
            2. Metodología (puntos principales)
            3. Hallazgos principales (máximo 5 puntos)
            4. Conclusiones clave (3-4 puntos)
            5. Recomendaciones (si existen)

            Usa viñetas y sé conciso pero informativo.
            """

            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": self.contexto_sistema},
                    {"role": "user", "content": prompt_inicial}
                ],
                temperature=0.7,
                max_tokens=1500
            )

            if progress_callback:
                progress_callback(0.6, "Refinando el resumen...")

            resumen = response.choices[0].message.content

            if progress_callback:
                progress_callback(1.0, "Resumen completado")

            return resumen

        except Exception as e:
            st.error(f"Error al generar el resumen: {str(e)}")
            return None

    def generar_respuesta(self, pregunta, texto_documento, nombre_documento, progress_callback=None):
        """Genera una respuesta a una pregunta específica"""
        try:
            if progress_callback:
                progress_callback(0.3, "Analizando la pregunta...")

            # Extraer contexto relevante
            contexto = texto_documento[:8000]  # Limitar el contexto inicial

            if progress_callback:
                progress_callback(0.5, "Generando respuesta...")

            prompt = f"""
            Basándote en el siguiente extracto del documento '{nombre_documento}':

            {contexto}

            Responde a esta pregunta: {pregunta}

            Proporciona una respuesta:
            1. Directa y específica
            2. Basada en el contenido del documento
            3. Con datos cuantitativos si están disponibles
            4. Citando secciones relevantes del documento
            """

            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": self.contexto_sistema},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1500
            )

            if progress_callback:
                progress_callback(1.0, "Respuesta completada")

            return response.choices[0].message.content

        except Exception as e:
            st.error(f"Error al generar la respuesta: {str(e)}")
            return None

    def crear_documento_respuesta(self, pregunta, respuesta, nombre_documento):
        """Crea un documento Word con la respuesta"""
        doc = Document()

        # Configurar estilos
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Título
        doc.add_heading('Análisis de Documento Energético', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Información general
        doc.add_heading('Información General', 1)
        p = doc.add_paragraph()
        p.add_run('Documento analizado: ').bold = True
        p.add_run(nombre_documento)
        p = doc.add_paragraph()
        p.add_run('Fecha de consulta: ').bold = True
        p.add_run(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

        # Pregunta y respuesta
        doc.add_heading('Consulta', 1)
        doc.add_paragraph(pregunta)

        doc.add_heading('Respuesta', 1)
        doc.add_paragraph(respuesta)

        # Guardar en memoria
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream

def main():
    st.title("Asistente Energético GPT")
    st.markdown("---")

    # Configuración en el sidebar
    st.sidebar.title("Configuración")
    api_key = st.sidebar.text_input("API Key de OpenAI", type="password")

    if not api_key:
        st.warning("Por favor, ingresa tu API key de OpenAI para continuar.")
        return

    # Inicializar asistente
    asistente = AsistenteEnergetico(api_key)

    # Área principal
    col1, col2 = st.columns([1, 2])

    with col1:
        st.subheader("Cargar Documento")
        archivo_subido = st.file_uploader("Selecciona un archivo PDF", type=['pdf'])

    if archivo_subido is not None:
        # Verificar tamaño del archivo
        if archivo_subido.size > 10 * 1024 * 1024:  # 10 MB
            st.error("El archivo es demasiado grande. Por favor, usa un PDF de menos de 10 MB.")
            return

        # Procesar PDF
        texto_documento = asistente.procesar_pdf(archivo_subido)

        if not texto_documento:
            st.error("No se pudo extraer texto del documento. Verifica que el PDF sea válido.")
            return

        with col2:
            st.subheader("Realizar Consulta")

            tipo_consulta = st.radio(
                "Tipo de consulta:",
                ["Resumen del documento", "Pregunta específica"]
            )

            if tipo_consulta == "Pregunta específica":
                pregunta = st.text_area("Escribe tu pregunta:", height=100)
            else:
                pregunta = "Genera un resumen completo del documento."

            if st.button("Generar respuesta"):
                progress_bar = st.progress(0)
                status_text = st.empty()

                def update_progress(progress, text):
                    progress_bar.progress(progress)
                    status_text.text(text)

                try:
                    with st.spinner("Procesando..."):
                        if tipo_consulta == "Resumen del documento":
                            respuesta = asistente.generar_resumen(
                                texto_documento,
                                archivo_subido.name,
                                update_progress
                            )
                        else:
                            if not pregunta.strip():
                                st.error("Por favor, ingresa una pregunta.")
                                return

                            respuesta = asistente.generar_respuesta(
                                pregunta,
                                texto_documento,
                                archivo_subido.name,
                                update_progress
                            )

                        if respuesta:
                            st.markdown("### Respuesta:")
                            st.markdown(respuesta)

                            # Generar documento
                            doc_stream = asistente.crear_documento_respuesta(
                                pregunta,
                                respuesta,
                                archivo_subido.name
                            )

                            # Botón de descarga
                            st.download_button(
                                label="Descargar respuesta como documento Word",
                                data=doc_stream.getvalue(),
                                file_name=f'Respuesta_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

                except Exception as e:
                    st.error(f"Error inesperado: {str(e)}")
                    st.info("Por favor, intenta nuevamente o contacta al soporte.")

if __name__ == "__main__":
    main()